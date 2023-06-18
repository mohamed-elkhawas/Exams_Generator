from PyQt5 import QtCore, QtWidgets
import sys
import pickle

from docx import Document
import random
import os

class Generator:
    def __init__(self):
        pass
    def work(self,subject="english",path="",no_type1=10,no_type2=10,no_exams=1):
        z = ["a", "b", "c", "d"]
        for i in range(no_exams):
            q1, q2 = [], []
            if path !="" :
                q1, q2 = self._read_the_questions(path)
                if no_type1 <= len(q1) and no_type2 <= len(q2):
                    n1 = self._generate_diff_random_numbers(no_questions=len(q1),wanted_num=no_type1)
                    n2 = self._generate_diff_random_numbers(no_questions=len(q2), wanted_num=no_type2)

                    pwd = os.getcwd()

                    dir = pwd + f"/{subject}"
                    if not os.path.exists(dir):
                        os.mkdir(dir)
                    dir = pwd + f"/{subject}/exams"
                    if not os.path.exists(dir):
                        os.mkdir(dir)
                    dir = pwd + f"/{subject}/model answer"
                    if not os.path.exists(dir):
                        os.mkdir(dir)


                    q_no = 1
                    document = Document()
                    document.add_heading(f'{subject} exams model {i+1}', 0)
                    document.add_heading('choose questions\n', 1)
                    for num in n1:
                        document.add_paragraph(f'{q_no}) {q1[num]["question"]}')
                        q_no += 1
                        for j in range(4):
                            document.add_paragraph(z[j]+"- "+q1[num]["choices"][j])

                    document.add_page_break()

                    document.add_heading('other questions\n', 1)

                    for num in n2:
                        if q2[num][1] != "":
                            document.add_paragraph(q2[num][1])
                        document.add_paragraph(f'{q_no}) {q2[num][0]["question"]}')
                        q_no += 1
                        for j in range(4):
                            document.add_paragraph(z[j]+"- "+q2[num][0]["choices"][j])

                    document.save(f'{subject}/exams/model {i+1}.docx')

                    q_no = 1
                    document = Document()
                    document.add_heading(f'{subject} model answer for model {i+1}', 0)

                    document.add_heading('choose questions', 1)
                    for num in n1:
                        document.add_paragraph(f'{q_no}) {q1[num]["answer"]}')
                        q_no += 1

                    document.add_heading('other questions\n', 1)
                    for num in n2:
                        document.add_paragraph(f'{q_no}) {q2[num][0]["answer"]}')
                        q_no += 1

                    document.save(f'{subject}/model answer/model {i+1}.docx')

    def _generate_diff_random_numbers(self, no_questions=10, wanted_num=10):
        random_nums, i = [], 0
        if wanted_num <= no_questions:
            while i < wanted_num:
                new_num = int(random.random() * no_questions)
                if new_num in random_nums:
                    pass
                else:
                    random_nums.append(new_num)
                    i += 1
        return random_nums

    def _read_the_questions(self, doc_path=r"Question bank.docx"):
        all_paras = Document(doc_path).paragraphs
        questions1, questions2, type_2_index = [], [], 0

        for i in range(len(all_paras)):
            try:
                question = {"question": "", "choices": [], "answer": ""}
                index = all_paras[i].text.index(":")
                if (int(all_paras[i].text[:index])):
                    question["question"] = all_paras[i - 5].text
                    question["answer"] = all_paras[i].text[index + 2:]
                    for j in range(4):
                        question["choices"].append(all_paras[i - 4 + j].text)

                    try:
                        idx = question["question"].index("…")
                        questions1.append(question)
                    except:
                        try:
                            idx = question["question"].index("...")
                            questions1.append(question)
                        except:
                            try:
                                idx = question["question"].index("passage")
                                question["question"] = all_paras[i - 5].text
                                questions2.append([question,all_paras[i - 6].text])
                            except:
                                if question["question"] != "":
                                    questions2.append([question,""])
            except:
                pass
        return questions1, questions2  # complete , passage

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.setEnabled(True)
        MainWindow.resize(587, 552)
        MainWindow.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.853, y1:0.909091, x2:0.021, y2:0.0340909, stop:0.0105263 rgba(0, 0, 0, 255), stop:1 rgba(102, 102, 102, 255))")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.startButton = QtWidgets.QPushButton(self.centralwidget)
        self.startButton.setGeometry(QtCore.QRect(240, 340, 93, 28))
        self.startButton.setStyleSheet("background-color:rgb(0, 85, 255);\n"
"color:white;")
        self.startButton.setObjectName("startButton")
        self.comboBox = QtWidgets.QComboBox(self.centralwidget)
        self.comboBox.setGeometry(QtCore.QRect(230, 270, 131, 26))
        self.comboBox.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.comboBox.setObjectName("comboBox")

        self.comboBox.clear()
        choices = []
        for choice in saved_choices["subjects"]:
            choices.append(choice)
            self.comboBox.addItems(choices)

        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(220, 170, 55, 16))
        self.label.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(30, 220, 121, 20))
        self.label_2.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(370, 220, 141, 21))
        self.label_3.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(170, 270, 55, 16))
        self.label_4.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(230, 420, 110, 50))
        self.label_5.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label_5.setAlignment(QtCore.Qt.AlignCenter)
        self.label_5.setIndent(-1)
        self.label_5.setObjectName("label_5")
        self.textEdit_4 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_4.setGeometry(QtCore.QRect(440, 40, 121, 31))
        self.textEdit_4.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_4.setObjectName("textEdit_4")
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setGeometry(QtCore.QRect(320, 40, 111, 30))
        self.label_6.setStyleSheet("font: 8pt \"MS Shell Dlg 2\";\n"
"color: white;\n"
"background-color:none;")
        self.label_6.setObjectName("label_6")


        self.browseButton = QtWidgets.QPushButton(self.centralwidget)
        self.browseButton.setGeometry(QtCore.QRect(20, 40, 221, 30))
        self.browseButton.setStyleSheet("background-color:rgb(0, 85, 255);\n"
"color:white;")
        self.browseButton.setObjectName("browseButton")
        self.addButton = QtWidgets.QPushButton(self.centralwidget)
        self.addButton.setGeometry(QtCore.QRect(210, 90, 141, 30))
        self.addButton.setStyleSheet("background-color:rgb(0, 85, 255);\n"
"color:white;")
        self.addButton.setObjectName("addButton")
        self.textEdit_5 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_5.setGeometry(QtCore.QRect(510, 220, 51, 26))
        self.textEdit_5.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_5.setObjectName("textEdit_5")
        self.textEdit_6 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_6.setGeometry(QtCore.QRect(290, 170, 51, 26))
        self.textEdit_6.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_6.setObjectName("textEdit_6")
        self.textEdit_7 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_7.setGeometry(QtCore.QRect(160, 220, 51, 26))
        self.textEdit_7.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.textEdit_7.setObjectName("textEdit_7")

#################### call my buttons fun #############

        self.choosed_path = ""
        self.choosed_sub = ""


        self.browseButton.clicked.connect(self._open_file_dialog)
        self.addButton.clicked.connect(self._update_combo_box)
        self.startButton.clicked.connect(self._start)
        self.startButton.clicked.connect(generator.work)

        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 587, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Exams Generator"))
        self.startButton.setText(_translate("MainWindow", "start"))
        self.label.setText(_translate("MainWindow", "no. exams"))
        self.label_2.setText(_translate("MainWindow", "no. choose Questions"))
        self.label_3.setText(_translate("MainWindow", "no. other Questions"))
        self.label_4.setText(_translate("MainWindow", "subject"))
        self.label_5.setText(_translate("MainWindow", "MADE BY MKK"))
        self.label_6.setText(_translate("MainWindow", "new subject name"))
        self.browseButton.setText(_translate("MainWindow", "browse new subject questions bank"))
        self.addButton.setText(_translate("MainWindow", "add the new subject"))

    def _open_file_dialog(self):  # a function to open the dialog window
        default_path =""
        self.choosed_path = QtWidgets.QFileDialog.getOpenFileName(None, "Open " + default_path + "MS Word files (*.docx)", '.', "(*.docx)")[0]

    def _update_combo_box(self):

        new_sub = self.textEdit_4.toPlainText()

        if new_sub != "" and self.choosed_path != "":
            saved_choices["subjects"][new_sub] = self.choosed_path

            self.comboBox.clear()
            choices = []
            for choice in saved_choices["subjects"]:
                choices.append(choice)
            self.comboBox.addItems(choices)

    def _start(self):
        no_1, no_2, no_exams = 0,0,0
        the_sub = ""
        try:
            the_sub = str(self.comboBox.currentText())
            no_exams = int(self.textEdit_6.toPlainText())
            no_1 = int(self.textEdit_7.toPlainText())
            no_2 = int(self.textEdit_5.toPlainText())
            if the_sub != "" :
                # for i in range(no_exams):
                generator.work(
                    subject=the_sub, path=saved_choices["subjects"][the_sub],
                    no_type1=no_1, no_type2=no_2, no_exams=no_exams
                )
                    # self.label_5.setText(QtCore.QCoreApplication.translate("MainWindow", f"finished {i+1} exams"))
                self.label_5.setText(QtCore.QCoreApplication.translate("MainWindow", f"finished ;)"))
                pickle.dump(saved_choices,open(saved_choices_path, 'wb'))
        except:
            self.label_5.setText(QtCore.QCoreApplication.translate("MainWindow", "try to close \nthe questions bank\n and try again"))
            print("passed")
            print(f"""path ={saved_choices["subjects"][the_sub]}  no_1 ={no_1}  no_2 ={no_2}    no_exams ={no_exams}    sub = {the_sub}""")

if __name__ == "__main__":

    saved_choices_path = r"don't_touch.pk1"
    try:
        saved_choices = pickle.load(open(saved_choices_path, 'rb'))
    except:
        saved_choices = {"subjects": {}, "no": {"exams": "1", "type1": "3", "type2": "3"}}

    generator = Generator()

    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
